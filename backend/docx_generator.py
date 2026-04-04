"""
DOCX 기획서 자동 생성
- 5개 섹션: 시장 분석 → 가격대 분포 → 진출 전략 → KPI → USP&인증
- 스타일: #2E75B6 헤더, Arial 폰트, 페이지 여백 1인치
- 스트리밍 지원
"""

import re
import json
from io import BytesIO
from datetime import datetime
from typing import List, Generator, Optional, Tuple
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from models import Product, Competitor, SessionMeta


# === 스타일 상수 ===

HEADER_COLOR = RGBColor(46, 117, 182)  # #2E75B6
LIGHT_HEADER_COLOR = RGBColor(235, 243, 251)  # #EBF3FB
FONT_NAME = "Arial"
HEADER1_SIZE = 16
HEADER2_SIZE = 14
BODY_TEXT_SIZE = 11


# === 문서 생성 함수 ===

def create_docx_report(
    meta: SessionMeta,
    products: List[Product],
    market_analysis: str,
    competitors: List[Competitor],
    competitor_analysis: str
) -> Generator[str, None, None]:
    """
    DOCX 기획서 생성 (스트리밍)

    Args:
        meta: 세션 메타데이터
        products: 상품 배열
        market_analysis: Step 2 분석 결과
        competitors: 경쟁사 배열
        competitor_analysis: Step 3 분석 결과

    Yields:
        진행 상태 메시지
    """
    # Document 생성
    doc = Document()

    # 페이지 설정: 상하좌우 1인치 여백
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # 제목 페이지
    title = doc.add_paragraph()
    title_run = title.add_run(f"{meta.brand_name} - {meta.category_name} 신제품 기획서")
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = HEADER_COLOR
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # 빈 줄
    doc.add_paragraph(f"작성일: {datetime.now().strftime('%Y년 %m월 %d일')}").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # === Section 1: 시장 분석 ===
    yield "섹션 1 생성 중..."
    _add_section_1_market_analysis(doc, products, market_analysis)
    yield "✅ 섹션 1: 시장 분석 완료\n"

    # === Section 2: 가격대 분포 분석 ===
    yield "섹션 2 생성 중..."
    _add_section_2_price_distribution(doc, market_analysis)
    yield "✅ 섹션 2: 가격대 분포 분석 완료\n"

    # === Section 3: 신제품 진출 전략 ===
    yield "섹션 3 생성 중..."
    _add_section_3_strategy(doc, market_analysis, meta.brand_name)
    yield "✅ 섹션 3: 신제품 진출 전략 완료\n"

    # === Section 4: 매출 목표 및 KPI ===
    yield "섹션 4 생성 중..."
    _add_section_4_kpi(doc, market_analysis)
    yield "✅ 섹션 4: 매출 목표 및 KPI 완료\n"

    # === Section 5: USP 및 예정 인증 ===
    yield "섹션 5 생성 중..."
    _add_section_5_usp_certification(doc, competitor_analysis)
    yield "✅ 섹션 5: USP 및 인증 완료\n"

    # 메모리에 저장
    doc_bytes = BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)

    # 전역 변수에 저장 (다운로드용)
    global _last_generated_doc
    _last_generated_doc = doc_bytes

    yield "기획서 생성 완료!"


# 전역 변수 (메모리 DOCX 저장용)
_last_generated_doc = None


# === 섹션별 작성 함수 ===

def _add_section_1_market_analysis(
    doc: Document,
    products: List[Product],
    market_analysis: str
) -> None:
    """Section 1: 시장 분석"""
    heading = doc.add_heading("1. 시장 분석", level=1)
    _style_heading(heading, size=HEADER1_SIZE)

    # TOP 25 테이블
    _style_heading(doc.add_heading("TOP 25 상품 정보", level=2))

    # 테이블 헤더
    table = doc.add_table(rows=1, cols=10)
    table.style = "Light Grid Accent 1"

    header_cells = table.rows[0].cells
    headers = ["순위", "상품명", "브랜드", "판매가", "100ml당", "리뷰수", "평점", "판매량", "형태", "추정매출"]

    for i, header_text in enumerate(headers):
        cell = header_cells[i]
        cell.text = header_text
        _style_header_cell(cell)

    # 데이터 행
    for product in products:
        row_cells = table.add_row().cells
        row_cells[0].text = str(product.rank or "")
        row_cells[1].text = product.name or ""
        row_cells[2].text = product.brand or ""
        row_cells[3].text = f"{product.price:,}원" if product.price else ""
        row_cells[4].text = f"{product.price_per_100ml:.0f}원" if product.price_per_100ml else ""
        row_cells[5].text = str(product.review_count or "")
        row_cells[6].text = f"{product.rating:.1f}" if product.rating else ""
        row_cells[7].text = product.sales_text or ""
        row_cells[8].text = product.form_type or ""
        row_cells[9].text = f"{product.revenue_estimate:,}원" if product.revenue_estimate else ""

        # 홀수 행 배경색
        if len(table.rows) % 2 == 0:
            for cell in row_cells:
                _style_data_cell(cell, light_bg=True)
        else:
            for cell in row_cells:
                _style_data_cell(cell)

    doc.add_paragraph()  # 빈 줄

    # 핵심 인사이트 추출
    _style_heading(doc.add_heading("핵심 인사이트", level=2))

    # 분석 결과에서 핵심 문구 추출
    insights = _extract_insights_from_analysis(market_analysis)
    for insight in insights:
        doc.add_paragraph(f'• {insight}')


def _add_section_2_price_distribution(
    doc: Document,
    market_analysis: str
) -> None:
    """Section 2: 가격대 분포 분석"""
    doc.add_page_break()
    heading = doc.add_heading("2. 가격대 분포 분석", level=1)
    _style_heading(heading)

    # 분석 결과에서 해당 섹션 추출
    section_text = _extract_section_text(market_analysis, 2)

    # 마크다운 테이블 파싱 및 변환
    _add_analysis_content(doc, section_text)


def _add_section_3_strategy(
    doc: Document,
    market_analysis: str,
    brand_name: str
) -> None:
    """Section 3: 신제품 진출 전략"""
    doc.add_page_break()
    heading = doc.add_heading("3. 신제품 진출 전략", level=1)
    _style_heading(heading)

    # Section 5 (신제품 스펙 제안) 추출
    section_text = _extract_section_text(market_analysis, 5)
    _add_analysis_content(doc, section_text)


def _add_section_4_kpi(
    doc: Document,
    market_analysis: str
) -> None:
    """Section 4: 매출 목표 및 KPI"""
    doc.add_page_break()
    heading = doc.add_heading("4. 매출 목표 및 KPI", level=1)
    _style_heading(heading)

    # 단계별 매출 목표 테이블 (예시 구조)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"

    header_cells = table.rows[0].cells
    header_cells[0].text = "단계"
    header_cells[1].text = "기간"
    header_cells[2].text = "월 매출 목표"

    for cell in header_cells:
        _style_header_cell(cell)

    # 샘플 데이터 (실제로는 Claude가 생성)
    stages = [
        ("Phase 1 (출시)", "1개월", "500만원"),
        ("Phase 2 (성장)", "2~4개월", "2,000만원"),
        ("Phase 3 (확대)", "5~12개월", "5,000만원"),
    ]

    for stage, period, target in stages:
        row_cells = table.add_row().cells
        row_cells[0].text = stage
        row_cells[1].text = period
        row_cells[2].text = target

        if stages.index((stage, period, target)) % 2 == 0:
            for cell in row_cells:
                _style_data_cell(cell, light_bg=True)
        else:
            for cell in row_cells:
                _style_data_cell(cell)

    doc.add_paragraph()
    _style_heading(doc.add_heading("핵심 KPI", level=2))

    kpis = [
        "순위 진입: 카테고리 TOP 20",
        "리뷰수 목표: 월 500+",
        "평점 유지: 4.5★ 이상",
    ]

    for kpi in kpis:
        doc.add_paragraph(f'• {kpi}')


def _add_section_5_usp_certification(
    doc: Document,
    competitor_analysis: str
) -> None:
    """Section 5: USP 및 예정 인증"""
    doc.add_page_break()
    heading = doc.add_heading("5. USP 및 예정 인증", level=1)
    _style_heading(heading)

    # 전체 분석 내용 추가
    _add_analysis_content(doc, competitor_analysis)


# === 스타일 함수 ===

def _style_heading(heading, color: RGBColor = HEADER_COLOR, size: int = None) -> None:
    """heading 단락의 runs[0]을 안전하게 스타일 적용 (runs 없을 때도 처리)"""
    if not heading.runs:
        run = heading.add_run()
    else:
        run = heading.runs[0]
    run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)


def _style_header_cell(cell) -> None:
    """테이블 헤더 셀 스타일"""
    # 배경색: #2E75B6
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), "2E75B6")
    cell._element.get_or_add_tcPr().append(shading_elm)

    # 텍스트 포맷
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)  # 흰 글씨
            run.font.bold = True
            run.font.name = FONT_NAME
            run.font.size = Pt(11)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _style_data_cell(cell, light_bg: bool = False) -> None:
    """테이블 데이터 셀 스타일"""
    if light_bg:
        # 배경색: #EBF3FB
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        shading_elm = OxmlElement("w:shd")
        shading_elm.set(qn("w:fill"), "EBF3FB")
        cell._element.get_or_add_tcPr().append(shading_elm)

    # 텍스트 포맷
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(11)


# === 보조 함수 ===

def _extract_section_text(analysis_text: str, section_num: int) -> str:
    """분석 결과에서 특정 섹션 추출"""
    pattern = rf"##\s*{section_num}[.\.]\s*(.+?)\n(.*?)(?=##|\Z)"
    match = re.search(pattern, analysis_text, re.DOTALL)

    if match:
        return match.group(0)
    return ""


def _extract_insights_from_analysis(analysis_text: str) -> List[str]:
    """분석 결과에서 핵심 인사이트 추출 (1~3줄)"""
    lines = analysis_text.split('\n')
    insights = []

    for i, line in enumerate(lines):
        line = line.strip()
        # 제목과 불릿 제외, 일반 텍스트만
        if line and not line.startswith("#") and not line.startswith("-") and not line.startswith("*"):
            if len(insights) < 3:
                insights.append(line)

    return insights[:3]


def _add_analysis_content(doc: Document, content: str) -> None:
    """분석 내용 추가 (텍스트 + 테이블)"""
    lines = content.split('\n')
    in_table = False
    table_lines = []

    for line in lines:
        line_stripped = line.strip()

        # 마크다운 테이블 감지
        if line_stripped.startswith("|") and "|" in line_stripped:
            if not in_table:
                in_table = True
            table_lines.append(line_stripped)
        else:
            if in_table:
                # 테이블 끝, 마크다운 테이블을 DOCX 테이블로 변환
                _add_markdown_table_to_doc(doc, table_lines)
                table_lines = []
                in_table = False

            # 일반 텍스트 추가
            if line_stripped:
                if line_stripped.startswith("###"):
                    # H3
                    heading = doc.add_heading(line_stripped.replace("### ", "").replace("###", ""), level=3)
                    _style_heading(heading)
                elif line_stripped.startswith("##"):
                    # H2
                    heading = doc.add_heading(line_stripped.replace("## ", "").replace("##", ""), level=2)
                    _style_heading(heading)
                elif line_stripped.startswith("-") or line_stripped.startswith("*"):
                    # 불릿
                    doc.add_paragraph(f'• {line_stripped[1:].strip()}')
                else:
                    # 일반 텍스트
                    doc.add_paragraph(line_stripped)

    # 마지막 테이블 처리
    if in_table and table_lines:
        _add_markdown_table_to_doc(doc, table_lines)


def _add_markdown_table_to_doc(doc: Document, table_lines: List[str]) -> None:
    """마크다운 테이블을 DOCX 테이블로 변환"""
    if len(table_lines) < 2:
        return

    # 헤더 파싱
    header_line = table_lines[0]
    headers = [h.strip() for h in header_line.split("|") if h.strip()]

    # 구분선 (테이블 형식 검증)
    # sep_line = table_lines[1]  # |---|---| 형식

    # 데이터 행 파싱
    data_lines = table_lines[2:] if len(table_lines) > 2 else []

    # DOCX 테이블 생성
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"

    # 헤더 스타일
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        _style_header_cell(cell)

    # 데이터 행 추가
    for data_line in data_lines:
        cells_data = [c.strip() for c in data_line.split("|") if c.strip()]
        if len(cells_data) == len(headers):
            row = table.add_row()
            for i, data in enumerate(cells_data):
                row.cells[i].text = data
                # 홀수 행 배경색
                if len(table.rows) % 2 == 0:
                    _style_data_cell(row.cells[i], light_bg=True)
                else:
                    _style_data_cell(row.cells[i])


def generate_file_name(meta: SessionMeta) -> str:
    """DOCX 파일명 생성"""
    today = datetime.now().strftime("%Y%m%d")
    return f"{meta.brand_name}_{meta.category_name}_기획서_{today}.docx"


def save_docx_to_bytes() -> BytesIO:
    """메모리의 DOCX를 바이트로 반환"""
    global _last_generated_doc
    if _last_generated_doc:
        _last_generated_doc.seek(0)
        return _last_generated_doc
    return BytesIO()
